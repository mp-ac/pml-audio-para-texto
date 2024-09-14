import torch
from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline
from docx import Document
import os
import gc
from pyannote.audio import Pipeline


class Transcrever:
    def __init__(self):
        self.diarization_pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization")

    def seconds_to_hms(self, seconds):
        """
        Converte segundos para horas, minutos e segundos
        :param seconds: int ou float
        :return: str
        """
        if seconds is None or not isinstance(seconds, (int, float)):
            return "---"
        else:
            hours = int(seconds // 3600)
            minutes = int((seconds % 3600) // 60)
            seconds = int(seconds % 60)
            return f"{hours:02}:{minutes:02}:{seconds:02}"

    def gerar_chunks(self, res, document, diarization_result):
        """
        Adiciona os chunks ao documento
        :param res: dict
        :param document: Document
        """
        # Verifica se a resposta é uma lista
        if isinstance(res, list):
            for item in res:
                if isinstance(item, dict) and 'chunks' in item:
                    for chunk in item['chunks']:
                        self.adicionar_paragrafo(chunk, document, diarization_result)
                elif isinstance(item, dict) and 'text' in item:
                    document.add_paragraph(item['text'])
        # Retorno quando a resposta contem timestamps
        elif isinstance(res, dict) and 'chunks' in res:
            for chunk in res['chunks']:
                self.adicionar_paragrafo(chunk, document, diarization_result)
        # Retorno quando a resposta é um texto sem timestamps
        elif isinstance(res, dict) and 'text' in res:
            document.add_paragraph(res['text'].replace('. ', '.\n\n'))
        else:
            raise TypeError("Formato de resultado inesperado.")

    def identificar_falante(self, timestamp, diarization_result):
        """
        Identifica o falante correspondente ao intervalo de tempo
        :param timestamp: tuple (start, end)
        :param diarization_result: Resultado da diarização
        :return: str
        """
        for turn, _, speaker in diarization_result.itertracks(yield_label=True):
            # Verifica se o intervalo do chunk coincide com o intervalo do falante
            if timestamp[0] >= turn.start and timestamp[1] <= turn.end:
                return f"Speaker {speaker}"
        return "Unknown Speaker"

    def adicionar_paragrafo(self, chunk, document, diarization_result):
        """
        Adiciona um parágrafo ao documento
        :param chunk: dict
        :param document: Document
        """
        if isinstance(chunk, dict) and 'timestamp' in chunk:
            start_time = self.seconds_to_hms(chunk['timestamp'][0])
            end_time = self.seconds_to_hms(chunk['timestamp'][1])

            # Identifica o falante no intervalo de tempo do chunk
            speaker = self.identificar_falante(chunk['timestamp'], diarization_result)

            input_dictionary = f"[{start_time} / {end_time}] - {speaker}: {chunk['text']}"
            document.add_paragraph(input_dictionary)
        else:
            raise TypeError("Formato de chunk inesperado.")

    def model(self):
        """
        Carrega o modelo Whisper v3 LARGE
        :return: pipeline
        """
        device = "cuda:0" if torch.cuda.is_available() else "cpu"
        torch_dtype = torch.float16 if torch.cuda.is_available() else torch.float32
        model_id = "openai/whisper-large-v3"

        model = AutoModelForSpeechSeq2Seq.from_pretrained(
                model_id,
                torch_dtype=torch_dtype,
                low_cpu_mem_usage=True,
                use_safetensors=True
        )
        model.to(device)

        processor = AutoProcessor.from_pretrained(model_id)

        pipe = pipeline(
            "automatic-speech-recognition",
            model=model,
            tokenizer=processor.tokenizer,
            feature_extractor=processor.feature_extractor,
            max_new_tokens=128,
            chunk_length_s=30,
            batch_size=16,
            return_timestamps=True,
            torch_dtype=torch_dtype,
            generate_kwargs={"language": "portuguese"},
            device=device,
        )

        return pipe

    def transcrever_audio(self, nome_arquivo, timestamp=True):
        """
        Inicia a transcrição do áudio
        :param nome_arquivo: str
        :param timestamp: bool - Adiciona ou não timestamps
        :return: str
        """
        diretorio_audios = 'audios'
        diretorio_transcritos = 'audios-transcritos'
        formato_arquivo_saida = '.docx'

        itens = os.listdir(diretorio_transcritos)

        arquivos = [item for item in itens if os.path.isfile(os.path.join(diretorio_transcritos, item))]

        quantidade_arquivos = len(arquivos) + 1

        if not os.path.exists(diretorio_transcritos):
            os.makedirs(diretorio_transcritos)

        nome_arquivo, extensao_arquivo = os.path.splitext(nome_arquivo)
        arquivo_docx = "transcricao_"+str(quantidade_arquivos)+"_"+nome_arquivo+formato_arquivo_saida

        caminho_completo_docx = os.path.join(diretorio_transcritos, arquivo_docx)

        if os.path.isfile(caminho_completo_docx):
            return caminho_completo_docx
        else:
            pipe = self.model()

            res = pipe(
                os.path.join(diretorio_audios, nome_arquivo + extensao_arquivo),
                return_timestamps=timestamp,
                chunk_length_s=30,
                stride_length_s=(5, 5)
            )

            # Realiza a diarização
            diarization_result = self.diarization_pipeline(
                    os.path.join(diretorio_audios, nome_arquivo + extensao_arquivo)
            )

            document = Document()

            self.gerar_chunks(res, document, diarization_result)

            document.save(caminho_completo_docx)

            # Libera memória e recursos
            del pipe
            gc.collect()
            torch.cuda.empty_cache()

            return caminho_completo_docx
